"""Build the Elysium data-provenance memo as a polished DOCX for PDF export."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).parent
SOURCE = ROOT / "DATA_PROVENANCE_MEMO.md"
OUTPUT = ROOT / "Elysium_Data_Provenance_Memo.docx"

NAVY = RGBColor(0x21, 0x3F, 0x57)
AQUA = RGBColor(0x31, 0xD5, 0xD0)
INK = RGBColor(0x17, 0x2C, 0x3D)
MUTED = RGBColor(0x6D, 0x82, 0x90)
LIGHT = "EDF8F8"


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=130, start=150, bottom=130, end=150) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "168F8C")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline_markdown(paragraph, text: str) -> None:
    pattern = re.compile(r"\[([^\]]+)\]\(([^)]+)\)|\*\*([^*]+)\*\*|`([^`]+)`")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        if match.group(1):
            add_hyperlink(paragraph, match.group(1), match.group(2))
        elif match.group(3):
            run = paragraph.add_run(match.group(3))
            run.bold = True
        elif match.group(4):
            run = paragraph.add_run(match.group(4))
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 14, 6),
        ("Heading 2", 12.5, NAVY, 10, 4),
        ("Heading 3", 11.5, NAVY, 8, 3),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = doc.styles["List Bullet"]
    bullet.font.name = "Aptos"
    bullet.font.size = Pt(10.2)
    bullet.font.color.rgb = INK
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.1


def add_running_furniture(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("ELYSIUM WEALTH MANAGEMENT  |  DATA PROVENANCE")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = MUTED

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Confidential business memorandum  |  McCade Pearson")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def add_masthead(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(3)
    run = kicker.add_run("DATA GOVERNANCE MEMORANDUM")
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = AQUA

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Traffic Accident Intelligence")
    run.font.name = "Aptos Display"
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = NAVY

    table = doc.add_table(rows=4, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(0.75)
    table.columns[1].width = Inches(5.8)
    details = [
        ("TO", "Leadership Team, Elysium Wealth Management"),
        ("FROM", "McCade Pearson  |  mccade.pearson@gmail.com"),
        ("DATE", "June 13, 2026"),
        ("SUBJECT", "Data sources, methodology, limitations, and appropriate use"),
    ]
    for row, (label, value) in zip(table.rows, details):
        row.cells[0].width = Inches(0.75)
        row.cells[1].width = Inches(5.8)
        for cell in row.cells:
            set_cell_margins(cell, top=55, bottom=55)
        lp = row.cells[0].paragraphs[0]
        lr = lp.add_run(label)
        lr.font.name = "Aptos"
        lr.font.size = Pt(8)
        lr.font.bold = True
        lr.font.color.rgb = MUTED
        vp = row.cells[1].paragraphs[0]
        vr = vp.add_run(value)
        vr.font.name = "Aptos"
        vr.font.size = Pt(9.5)
        vr.font.color.rgb = INK

    callout = doc.add_table(rows=1, cols=1)
    callout.autofit = False
    cell = callout.cell(0, 0)
    set_cell_fill(cell, LIGHT)
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(
        "Purpose: document the authoritative government sources and analytical methods "
        "used by Elysium's Traffic Accident Intelligence application."
    )
    run.font.name = "Aptos"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = NAVY


def build() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style_document(doc)
    add_running_furniture(doc)
    add_masthead(doc)

    skip_title_block = True
    in_code = False
    code_lines: list[str] = []
    for line in lines:
        stripped = line.strip()
        if skip_title_block:
            if stripped == "## Purpose":
                skip_title_block = False
            else:
                continue

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = NAVY
                code_lines = []
            in_code = not in_code
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not stripped:
            continue
        if stripped.startswith("## "):
            doc.add_paragraph(stripped[3:], style="Heading 1")
            continue
        if stripped.startswith("### "):
            doc.add_paragraph(stripped[4:], style="Heading 2")
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, stripped[2:])
            continue
        if stripped.startswith("**McCade Pearson**"):
            p = doc.add_paragraph()
            r = p.add_run("McCade Pearson")
            r.bold = True
            r.font.color.rgb = NAVY
            continue
        p = doc.add_paragraph()
        add_inline_markdown(p, stripped)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()

